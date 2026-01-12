# utils.py
from PyPDF2 import PdfReader
from docx import Document
import io
import logging
from fpdf import FPDF
import re
from urllib.parse import urlparse, urlunparse

class Struct:
    """Convert dict to object for config"""
    def __init__(self, **entries):
        self.__dict__.update(entries)

def canonicalize_url(url):
    if not url:
        return url
    # Remove query parameters and fragments for canonical duplicate check
    parts = urlparse(url)
    clean_url = urlunparse((parts.scheme, parts.netloc, parts.path, '', '', ''))
    return clean_url.rstrip('/')

def extract_text_from_file(file):
    """Extract text from PDF, DOCX, or TXT files with error handling"""
    try:
        content = []
        file_bytes = file.getvalue()

        if file.type == "application/pdf":
            pdf = PdfReader(io.BytesIO(file_bytes))
            for page in pdf.pages:
                if text := page.extract_text():
                    try:
                        content.append(text.encode('utf-8').decode('utf-8'))
                    except:
                        content.append(text.encode('latin-1', 'replace').decode('latin-1'))
            # Fallback for image-based PDFs
            if not content:
                logging.warning("PDF appears to be image-based - text extraction limited")
                content = ["[PDF content requires OCR extraction]"]

        elif file.type == "text/plain":
            # Handle different encodings
            try:
                text = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                text = file_bytes.decode('latin-1')
            content = [text]

        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(io.BytesIO(file_bytes))
            doc_content = []

            # Process paragraphs with proper spacing
            for para in doc.paragraphs:
                if para.text.strip():
                    doc_content.append(para.text + "\n")

            # Process tables with cell separation
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = ' '.join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                        row_data.append(cell_text)
                    doc_content.append(" | ".join(row_data))
                    if row_idx == 0:  # Add header separator
                        doc_content.append("-" * 50)

            content = doc_content  # Assign to main content list

        full_content = "\n".join(content).strip()
        full_content = re.sub(r'\n{3,}', '\n\n', full_content)
        logging.info(f"Extracted {len(full_content)} characters")
        return full_content

    except Exception as e:
        logging.error(f"Text extraction failed: {str(e)}")
        return ""

def create_pdf(content, filename="customized_cv.pdf"):
    """Create a styled PDF document from text content - matches source CV style"""
    try:
        pdf = FPDF()
        pdf.set_margins(left=15, top=12, right=15)
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=12)

        # Font configuration - including italic for styled text
        pdf.add_font('Inter', '', 'fonts/Inter_18pt-Regular.ttf', uni=True)
        pdf.add_font('Inter', 'B', 'fonts/Inter_18pt-Bold.ttf', uni=True)
        pdf.add_font('Inter', 'I', 'fonts/Inter_18pt-Italic.ttf', uni=True)
        font_name = 'Inter'
        base_font_size = 10  # Slightly smaller for compact look
        line_height = 5  # Tighter line spacing

        # Process content
        cleaned_content = re.sub(r'\n{3,}', '\n\n', content.strip())
        current_section = None
        name_printed = False

        for line in cleaned_content.split('\n'):
            line = line.strip()
            if not line:
                continue

            # Title Section (Name) - UPPERCASE style
            if line.startswith('# '):
                name = line[2:].strip().upper()
                _handle_title_compact(pdf, name, font_name)
                name_printed = True
                continue

            # Single-line contact info with pipe separators (right after name)
            if name_printed and '|' in line and not line.startswith('**') and not line.startswith('##'):
                _handle_contact_line(pdf, line, font_name, base_font_size, line_height)
                name_printed = False  # Reset after contact line
                continue

            # Section Headers
            if line.startswith('## '):
                current_section = line[3:].strip().upper()
                _handle_section_header_compact(pdf, current_section, font_name)
                continue

            # Employment History - Single line format with pipes
            if line.startswith('**') and '|' in line:
                _handle_employment_entry_compact(pdf, line, font_name, base_font_size, line_height)
                continue

            # Sub-headers (bold text without pipes)
            if line.startswith('**') and not '|' in line:
                _handle_subheader(pdf, line, font_name, base_font_size)
                continue

            # Handle bullet points (both • and - styles)
            if line.startswith('•') or line.startswith('-'):
                bullet_char = line[0]
                is_tech = '**' in line and ':' in line
                _handle_bullet_point_compact(pdf, line, font_name, base_font_size, line_height, is_tech, bullet_char)
                continue

            # Default text formatting (for paragraphs like Executive Profile)
            _handle_paragraph_text(pdf, line, font_name, base_font_size, line_height)

        pdf.output(filename)
        return filename

    except Exception as e:
        logging.error(f"PDF creation failed: {str(e)}")
        return None


def create_docx(content, filename="customized_cv.docx"):
    """Create a styled Word document from markdown content - matches source CV style"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Set narrow margins for compact layout
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

        # Process content
        cleaned_content = re.sub(r'\n{3,}', '\n\n', content.strip())
        name_printed = False

        for line in cleaned_content.split('\n'):
            line = line.strip()
            if not line:
                continue

            # Title Section (Name) - UPPERCASE style
            if line.startswith('# '):
                name = line[2:].strip().upper()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(name)
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Calibri'
                p.paragraph_format.space_after = Pt(2)
                name_printed = True
                continue

            # Single-line contact info with pipe separators
            if name_printed and '|' in line and not line.startswith('**') and not line.startswith('##'):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line)
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
                p.paragraph_format.space_after = Pt(6)
                name_printed = False
                continue

            # Section Headers with shading
            if line.startswith('## '):
                section_title = line[3:].strip().upper()
                p = doc.add_paragraph()
                _add_shading_to_paragraph(p)
                run = p.add_run(section_title)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
                continue

            # Employment History - Single line format with pipes
            if line.startswith('**') and '|' in line:
                entry = line.replace('**', '').strip()
                p = doc.add_paragraph()
                run = p.add_run(entry)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
                continue

            # Sub-headers (bold text without pipes)
            if line.startswith('**') and '|' not in line:
                subheader = line.replace('**', '').strip()
                p = doc.add_paragraph()
                run = p.add_run(subheader)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                p.paragraph_format.space_after = Pt(2)
                continue

            # Handle bullet points
            if line.startswith('•') or line.startswith('-'):
                bullet_text = line[1:].strip()
                p = doc.add_paragraph(style='List Bullet')
                _add_formatted_text_to_paragraph(p, bullet_text)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Inches(0.2)
                continue

            # Default text formatting (for paragraphs like Executive Profile)
            p = doc.add_paragraph()
            _add_formatted_text_to_paragraph(p, line)
            p.paragraph_format.space_after = Pt(4)

        doc.save(filename)
        return filename

    except Exception as e:
        logging.error(f"DOCX creation failed: {str(e)}")
        return None


def _add_shading_to_paragraph(paragraph):
    """Add light gray shading to a paragraph (for section headers)"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E6E6E6')  # Light gray (#E6E6E6)
    paragraph._p.get_or_add_pPr().append(shading_elm)


def _add_formatted_text_to_paragraph(paragraph, text):
    """Add text with markdown formatting support (bold and italic) to a Word paragraph"""
    from docx.shared import Pt

    # Split by bold markers first
    bold_segments = re.split(r'(\*\*)', text)
    is_bold = False

    for bold_seg in bold_segments:
        if not bold_seg:
            continue
        if bold_seg == '**':
            is_bold = not is_bold
            continue

        # Process italic within this segment
        italic_segments = re.split(r'(?<!\*)(\*)(?!\*)', bold_seg)
        is_italic = False

        for italic_seg in italic_segments:
            if not italic_seg:
                continue
            if italic_seg == '*':
                is_italic = not is_italic
                continue

            # Add run with appropriate formatting
            run = paragraph.add_run(italic_seg)
            run.bold = is_bold
            run.italic = is_italic
            run.font.size = Pt(10)
            run.font.name = 'Calibri'


def _handle_text(pdf, text, font_name, base_font_size, line_height):
    """Improved markdown handling with proper bold formatting"""
    # Split text into segments while preserving order
    segments = re.split(r'(\*\*)', text)
    is_bold = False
    x = pdf.get_x()
    y = pdf.get_y()

    for segment in segments:
        if not segment:
            continue

        if segment == '**':
            is_bold = not is_bold
            continue

        # Set appropriate font
        if is_bold:
            pdf.set_font(font_name, 'B', base_font_size)
        else:
            pdf.set_font(font_name, '', base_font_size)

        # Handle word wrapping
        words = segment.split()
        for word in words:
            word_width = pdf.get_string_width(word + ' ')

            # Check page space
            if x + word_width > pdf.w - pdf.r_margin:
                x = pdf.l_margin
                y += line_height
                if y > pdf.h - 15:
                    pdf.add_page()
                    y = pdf.t_margin

            pdf.set_xy(x, y)
            pdf.cell(word_width, line_height, word + ' ', ln=0)
            x += word_width

    # Move to next line after processing all segments
    pdf.ln(line_height)
    pdf.set_x(pdf.l_margin)


def _handle_bullet_point(pdf, line, font_name, base_font_size, line_height, is_tech_skills=False):
    """Consistent bullet alignment with proper multi-line wrapping and markdown support"""
    bullet_text = line[2:].strip()
    initial_x = pdf.get_x()

    if is_tech_skills and ':' in bullet_text:
        # Technology skills with category bullet
        category, items = re.split(r':\s*', bullet_text, 1)
        category = category.replace('**', '').strip()
        items = items.replace('**', '').strip()
        # Print bullet point, category in bold, and items in regular font
        pdf.set_font(font_name, '', base_font_size)
        pdf.cell(8, line_height, '•', ln=0)
        pdf.set_font(font_name, 'B', base_font_size)
        category_width = pdf.get_string_width(f"{category}: ")
        pdf.cell(category_width, line_height, f"{category}: ", ln=0)
        pdf.set_font(font_name, '', base_font_size)
        pdf.multi_cell(0, line_height, items, ln=1)
    else:
        # Regular bullet points with markdown support
        pdf.set_font(font_name, '', base_font_size)
        pdf.set_x(initial_x + 8)  # Indent for bullet
        pdf.cell(8, line_height, '•', ln=0)  # Print bullet

        # Save position after bullet
        start_x = pdf.get_x()
        start_y = pdf.get_y()

        # Process text with markdown handling
        segments = re.split(r'(\*\*)', bullet_text)
        is_bold = False
        pdf.set_xy(start_x, start_y)

        for segment in segments:
            if not segment:
                continue
            if segment == '**':
                is_bold = not is_bold
                continue

            # Set appropriate font style
            pdf.set_font(font_name, 'B' if is_bold else '', base_font_size)

            # Write segment content
            pdf.write(line_height, segment)

        # Move to next line after bullet content
        pdf.ln(line_height)
        pdf.set_x(initial_x + 8)  # Maintain bullet indentation

    pdf.ln(2)

def _handle_employment_entry(pdf, line, font_name, base_font_size, line_height):
    """Improved company/role layout with page break check"""
    entry = line.replace('**', '').strip()
    company, role = entry.split('|', 1)
    company = company.strip()
    role = role.strip()
    # Check available space
    if pdf.get_y() > pdf.h - 20:
        pdf.add_page()
    # Save original position
    start_x = pdf.get_x()
    start_y = pdf.get_y()
    # Print company name
    pdf.set_font(font_name, 'B', base_font_size)
    company_width = pdf.get_string_width(company) + 2
    pdf.cell(company_width, line_height, company, ln=0)
    # Calculate role position
    role_x = start_x + company_width + 2
    if role_x > pdf.w - 15:  # Prevent overflow
        role_x = start_x
        start_y += line_height
    # Print role with proper wrapping
    pdf.set_font(font_name, '', base_font_size - 1)
    pdf.set_xy(role_x, start_y)
    pdf.multi_cell(0, line_height, role)
    # Maintain vertical alignment
    new_y = max(start_y + line_height, pdf.get_y())
    pdf.set_xy(start_x, new_y)
    pdf.ln(2)

def _process_contact_section(pdf, contact_lines, font_name, font_size, line_height):
    """Properly centered contact information with vertical alignment"""
    # Calculate maximum width of any contact line
    pdf.set_font(font_name, 'B', font_size)
    max_width = max(pdf.get_string_width(line.replace('**', '')) for line in contact_lines)
    # Center all lines horizontally
    start_x = (pdf.w - max_width) / 2
    start_y = pdf.get_y()
    for line in contact_lines:
        pdf.set_xy(start_x, start_y)
        if ':' in line:
            label, value = line.split(':', 1)
            # Calculate combined width
            pdf.set_font(font_name, 'B', font_size)
            label_width = pdf.get_string_width(label + ': ')
            pdf.set_font(font_name, '', font_size)
            value_width = pdf.get_string_width(value.strip())
            total_width = label_width + value_width
            # Center this specific line
            line_x = (pdf.w - total_width) / 2
            pdf.set_x(line_x)
            pdf.set_font(font_name, 'B', font_size)
            pdf.cell(label_width, line_height, label + ': ', ln=0)
            pdf.set_font(font_name, '', font_size)
            pdf.cell(value_width, line_height, value.strip(), ln=1)
        else:
            pdf.set_x(start_x)
            pdf.cell(max_width, line_height, line, ln=1, align='C')
        start_y += line_height  # Maintain vertical spacing
    pdf.ln(4)

def _handle_title(pdf, title_text, font_name):
    """Format the main title section"""
    pdf.set_font(font_name, 'B', 16)
    pdf.cell(0, 8, title_text, 0, 1, 'C')
    pdf.ln(4)

def _handle_subheader(pdf, line, font_name, base_font_size):
    """Format subheaders in Key Achievements section"""
    subheader = line.replace('**', '').strip()
    pdf.set_font(font_name, 'B', base_font_size)
    pdf.multi_cell(0, base_font_size * 0.6, subheader)
    pdf.ln(2)

def _handle_section_header(pdf, title, font_name):
    """Format section headers with consistent spacing"""
    pdf.ln(4)  # Space before header
    pdf.set_font(font_name, 'B', 12)
    pdf.cell(0, 8, title.upper(), 0, 1, 'C')
    pdf.ln(4)  # Space after header


# ============ NEW COMPACT STYLE FUNCTIONS ============

def _handle_title_compact(pdf, name, font_name):
    """Format name in compact uppercase style - centered"""
    pdf.set_font(font_name, 'B', 14)
    pdf.cell(0, 6, name, 0, 1, 'C')
    pdf.ln(1)


def _handle_contact_line(pdf, line, font_name, base_font_size, line_height):
    """Format single-line contact info with pipe separators - centered"""
    pdf.set_font(font_name, '', base_font_size - 1)
    # Center the contact line
    contact_width = pdf.get_string_width(line)
    x_pos = (pdf.w - contact_width) / 2
    pdf.set_x(x_pos)
    pdf.cell(contact_width, line_height, line, 0, 1)
    pdf.ln(3)


def _handle_section_header_compact(pdf, title, font_name):
    """Format section headers - compact style with shaded background like source CV"""
    pdf.ln(4)  # Space before header

    # Save current position
    x = pdf.get_x()
    y = pdf.get_y()

    # Draw shaded background rectangle (light gray)
    pdf.set_fill_color(230, 230, 230)  # Light gray background
    header_height = 6
    pdf.rect(pdf.l_margin, y, pdf.w - pdf.l_margin - pdf.r_margin, header_height, 'F')

    # Set position and print header text
    pdf.set_xy(pdf.l_margin + 2, y + 1)  # Slight padding inside the box
    pdf.set_font(font_name, 'B', 10)
    pdf.set_text_color(0, 0, 0)  # Black text
    pdf.cell(0, 4, title.upper(), 0, 1, 'L')

    # Reset position after header
    pdf.set_y(y + header_height + 2)
    pdf.set_text_color(0, 0, 0)  # Reset to black


def _handle_employment_entry_compact(pdf, line, font_name, base_font_size, line_height):
    """Format employment entry - all on single line with pipes"""
    entry = line.replace('**', '').strip()
    # Check available space
    if pdf.get_y() > pdf.h - 15:
        pdf.add_page()
    pdf.set_font(font_name, 'B', base_font_size)
    pdf.multi_cell(0, line_height, entry)
    pdf.ln(1)


def _handle_bullet_point_compact(pdf, line, font_name, base_font_size, line_height, is_tech_skills=False, bullet_char='•'):
    """Compact bullet points with proper indentation and full markdown support"""
    # Remove the bullet character from the start
    bullet_text = line[1:].strip() if line[0] in ['•', '-'] else line.strip()
    indent = 5  # Bullet indentation

    if is_tech_skills and ':' in bullet_text:
        # Technology skills with category
        category, items = re.split(r':\s*', bullet_text, 1)
        category = category.replace('**', '').strip()
        items = items.replace('**', '').strip()

        pdf.set_x(pdf.l_margin + indent)
        pdf.set_font(font_name, '', base_font_size)
        pdf.cell(4, line_height, '•', ln=0)
        pdf.set_font(font_name, 'B', base_font_size)
        category_width = pdf.get_string_width(f"{category}: ")
        pdf.cell(category_width, line_height, f"{category}: ", ln=0)
        pdf.set_font(font_name, '', base_font_size)
        # Calculate remaining width for items
        remaining_width = pdf.w - pdf.get_x() - pdf.r_margin
        pdf.multi_cell(remaining_width, line_height, items)
    else:
        # Regular bullet points with full markdown support (bold + italic)
        pdf.set_x(pdf.l_margin + indent)
        pdf.set_font(font_name, '', base_font_size)
        pdf.cell(4, line_height, '•', ln=0)

        # Use the formatted text writer for full markdown support
        _write_formatted_text(pdf, bullet_text, font_name, base_font_size, line_height)

        pdf.ln(line_height)

    pdf.ln(0.5)  # Minimal spacing between bullets


def _handle_paragraph_text(pdf, text, font_name, base_font_size, line_height):
    """Handle paragraph text (like Executive Profile) with markdown formatting support"""
    # Process markdown formatting: **bold** and *italic*
    _write_formatted_text(pdf, text, font_name, base_font_size, line_height)
    pdf.ln(line_height + 1)


def _write_formatted_text(pdf, text, font_name, base_font_size, line_height):
    """Write text with markdown formatting support (bold and italic)"""
    # Split by bold markers first, then handle italic within each segment
    bold_segments = re.split(r'(\*\*)', text)
    is_bold = False

    for bold_seg in bold_segments:
        if not bold_seg:
            continue
        if bold_seg == '**':
            is_bold = not is_bold
            continue

        # Now process italic within this segment
        italic_segments = re.split(r'(?<!\*)(\*)(?!\*)', bold_seg)
        is_italic = False

        for italic_seg in italic_segments:
            if not italic_seg:
                continue
            if italic_seg == '*':
                is_italic = not is_italic
                continue

            # Determine font style
            if is_bold and is_italic:
                # Note: fpdf doesn't support BI style, use bold as priority
                pdf.set_font(font_name, 'B', base_font_size)
            elif is_bold:
                pdf.set_font(font_name, 'B', base_font_size)
            elif is_italic:
                pdf.set_font(font_name, 'I', base_font_size)
            else:
                pdf.set_font(font_name, '', base_font_size)

            pdf.write(line_height, italic_seg)



